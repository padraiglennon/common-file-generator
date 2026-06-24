"""Complexity-driven Word document generator.

Builds a ``.docx`` from scratch (not template injection). A document is a series
of **sections**; each section opens with a heading and a paragraph, then draws
extra content blocks from a weighted pool that grows with the complexity level.
Content is seeded lorem-ipsum and placeholder images are generated in memory, so
the same ``(complexity, sections, seed)`` always yields the identical document
and no external assets are needed.
"""

from __future__ import annotations

import io
import random
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from PIL import Image

from ms_office_file_generator.core.complexity import Complexity, doc_block_pool
from ms_office_file_generator.core.lorem import Lorem

_USABLE_WIDTH = Inches(6.0)

# Extra content blocks per section, by complexity. Richer levels are denser, so
# section richness scales with complexity (not just which block types appear).
_BLOCKS_PER_SECTION: dict[Complexity, int] = {
    Complexity.MINIMAL: 1,
    Complexity.SIMPLE: 2,
    Complexity.STANDARD: 3,
    Complexity.COMPLEX: 4,
    Complexity.MAXIMUM: 5,
}


class DocxComplexityGenerator:
    """Generate a Word document at a chosen complexity level."""

    def __init__(
        self,
        complexity: Complexity = Complexity.STANDARD,
        *,
        sections: int = 5,
        seed: int = 0,
        blocks_per_section: int | None = None,
    ) -> None:
        if sections < 1:
            raise ValueError("sections must be at least 1")
        if blocks_per_section is not None and blocks_per_section < 1:
            raise ValueError("blocks_per_section must be at least 1")
        self.complexity = complexity
        self.sections = sections
        self.blocks_per_section = (
            blocks_per_section
            if blocks_per_section is not None
            else _BLOCKS_PER_SECTION[complexity]
        )
        pool = doc_block_pool(complexity)
        self._block_keys = [key for key, _ in pool]
        self._block_weights = [weight for _, weight in pool]
        self._rng = random.Random(seed)
        self._lorem = Lorem(self._rng)

    def build(self) -> Document:
        document = Document()
        document.add_heading(self._lorem.title(6), level=0)  # document title
        for index in range(1, self.sections + 1):
            self._add_section(document, index)
        return document

    def save(self, out_path: str | Path) -> Path:
        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        self.build().save(str(out_path))
        return out_path

    # --- sections + blocks ----------------------------------------------

    def _add_section(self, document: Document, index: int) -> None:
        document.add_heading(f"{index}. {self._lorem.title()}", level=1)
        self._add_paragraph(document)
        for _ in range(self.blocks_per_section):
            key = self._rng.choices(self._block_keys, weights=self._block_weights)[0]
            self._BLOCKS[key](self, document)

    def _add_paragraph(self, document: Document) -> None:
        document.add_paragraph(self._lorem.paragraph(2, 4))

    def _add_bullets(self, document: Document) -> None:
        document.add_heading(self._lorem.title(3), level=2)
        for text in self._lorem.bullets(self._rng.randint(3, 5)):
            document.add_paragraph(text, style="List Bullet")

    def _add_numbered(self, document: Document) -> None:
        document.add_heading(self._lorem.title(3), level=2)
        for text in self._lorem.bullets(self._rng.randint(3, 5)):
            document.add_paragraph(text, style="List Number")

    def _add_table(self, document: Document) -> None:
        cols = self._rng.randint(3, 4)
        rows = self._rng.randint(3, 6)
        table = document.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.autofit = False
        width = Inches(_USABLE_WIDTH.inches / cols)
        for c in range(cols):
            table.cell(0, c).text = self._lorem.title(2)
        for r in range(1, rows):
            for c in range(cols):
                table.cell(r, c).text = self._lorem.words(self._rng.randint(1, 3))
        for row in table.rows:
            for cell in row.cells:
                cell.width = width

    def _add_image(self, document: Document) -> None:
        document.add_picture(self._placeholder_png(), width=Inches(4.5))

    def _add_quote(self, document: Document) -> None:
        document.add_paragraph(self._lorem.sentence(8, 16), style="Intense Quote")

    def _add_page_break(self, document: Document) -> None:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    _BLOCKS = {
        "paragraph": _add_paragraph,
        "bullets": _add_bullets,
        "numbered": _add_numbered,
        "table": _add_table,
        "image": _add_image,
        "quote": _add_quote,
        "page_break": _add_page_break,
    }

    # --- helpers --------------------------------------------------------

    def _placeholder_png(self) -> io.BytesIO:
        color = (
            self._rng.randint(40, 220),
            self._rng.randint(40, 220),
            self._rng.randint(40, 220),
        )
        image = Image.new("RGB", (480, 270), color)
        stream = io.BytesIO()
        image.save(stream, format="PNG")
        stream.seek(0)
        return stream
