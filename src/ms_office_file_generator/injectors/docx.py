"""Word injector (python-docx).

Word documents do not expose user-settable shape names the way PowerPoint does,
so named elements use in-document marker tokens instead:

* ``{{token}}``        - plain text placeholder, replaced with a value.
* ``{{table:NAME}}``   - placed in any cell of a table to give it the name NAME.
* ``{{media:NAME}}``   - placed in its own paragraph; replaced by an image.

This keeps the named-tag philosophy consistent across formats while working
within Word's model. No tables or paragraphs are created; only filled.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.table import Table

from ms_office_file_generator.core.config import Config, MediaSpec, TableSpec
from ms_office_file_generator.core.injector import Injector
from ms_office_file_generator.injectors._text import (
    find_tokens,
    replace_in_paragraph,
)

_TABLE_MARKER = re.compile(r"\{\{\s*table:([\w.-]+)\s*\}\}")
_MEDIA_MARKER = re.compile(r"\{\{\s*media:([\w.-]+)\s*\}\}")


class DocxInjector(Injector):
    """Inject data into a Word Golden template."""

    extensions = frozenset({".docx"})

    def _load(self, path: Path) -> Document:
        return Document(str(path))

    def _save(self, path: Path) -> None:
        self._document.save(str(path))

    def _body_paragraphs(self) -> list:
        return list(self._document.paragraphs)

    # --- text -----------------------------------------------------------

    def _inject_text(self, config: Config) -> None:
        if not config.text:
            return
        seen: set[str] = set()
        for paragraph in self._iter_all_paragraphs():
            seen |= find_tokens(paragraph.text)
            replace_in_paragraph(paragraph, config.text)

        for name in sorted(set(config.text) - seen):
            self.report.warning(
                f"text.{name}",
                f"no '{{{{{name}}}}}' placeholder was found in the template; "
                "this value was not used.",
            )

    def _iter_all_paragraphs(self):
        yield from self._document.paragraphs
        for table in self._document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    # --- tables ---------------------------------------------------------

    def _inject_tables(self, config: Config) -> None:
        if not config.tables:
            return
        named = self._named_tables()
        for spec in config.tables:
            table = named.get(spec.name)
            if table is None:
                self.report.error(
                    f"tables.{spec.name}",
                    f"no '{{{{table:{spec.name}}}}}' marker was found in any table.",
                )
                continue
            self._fill_table(table, spec)

    def _named_tables(self) -> dict[str, Table]:
        named: dict[str, Table] = {}
        for table in self._document.tables:
            for row in table.rows:
                for cell in row.cells:
                    match = _TABLE_MARKER.search(cell.text)
                    if match:
                        named.setdefault(match.group(1), table)
        return named

    def _fill_table(self, table: Table, spec: TableSpec) -> None:
        n_rows, n_cols = len(table.rows), len(table.columns)
        data_rows = list(spec.rows)
        if spec.header is not None:
            data_rows = [spec.header, *data_rows]

        if len(data_rows) > n_rows:
            self.report.warning(
                f"tables.{spec.name}",
                f"you provided {len(data_rows)} rows but the template table holds "
                f"{n_rows}; kept the first {n_rows}.",
            )
        for r, row in enumerate(data_rows[:n_rows]):
            if len(row) > n_cols:
                self.report.warning(
                    f"tables.{spec.name}",
                    f"row {r} has {len(row)} cells but the table has {n_cols} "
                    f"columns; kept the first {n_cols}.",
                )
            for c, value in enumerate(row[:n_cols]):
                table.cell(r, c).text = value

    # --- media ----------------------------------------------------------

    def _inject_media(self, config: Config) -> None:
        if not config.media:
            return
        by_name = {spec.name: spec for spec in config.media}
        placed: set[str] = set()
        for paragraph in self._iter_all_paragraphs():
            match = _MEDIA_MARKER.search(paragraph.text)
            if not match:
                continue
            name = match.group(1)
            spec = by_name.get(name)
            if spec is None:
                continue
            if self._place_media(paragraph, spec):
                placed.add(name)

        for name in sorted(set(by_name) - placed):
            self.report.error(
                f"media.{name}",
                f"no '{{{{media:{name}}}}}' marker was found in the template.",
            )

    def _place_media(self, paragraph, spec: MediaSpec) -> bool:
        if not spec.path.is_file():
            self.report.error(
                f"media.{spec.name}",
                f"image file not found: {spec.path}.",
            )
            return False
        for run in paragraph.runs:
            run.text = ""
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.add_picture(str(spec.path))
        return True
