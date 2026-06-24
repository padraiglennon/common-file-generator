"""Tests for the complexity-driven Word document generator."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from ms_office_file_generator.core import Complexity, generate_doc
from ms_office_file_generator.generators import DocxComplexityGenerator


def test_section_count_drives_headings(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="standard", sections=8, seed=1)
    doc = Document(str(out))
    h1 = [p for p in doc.paragraphs if p.style and p.style.name == "Heading 1"]
    assert len(h1) == 8


def test_invalid_section_count_raises() -> None:
    with pytest.raises(ValueError, match="at least 1"):
        DocxComplexityGenerator(Complexity.MINIMAL, sections=0)


def test_minimal_has_no_tables_or_images(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="minimal", sections=10, seed=3)
    doc = Document(str(out))
    assert len(doc.tables) == 0
    assert len(doc.inline_shapes) == 0


def test_maximum_has_tables_and_images(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="maximum", sections=25, seed=7)
    doc = Document(str(out))
    assert len(doc.tables) > 0
    assert len(doc.inline_shapes) > 0
    styles = {p.style.name for p in doc.paragraphs if p.style}
    assert "List Bullet" in styles or "List Number" in styles


def test_blocks_per_section_scales_with_complexity(tmp_path: Path) -> None:
    # Maximum is denser than minimal at the same section count, so it has more
    # body paragraphs/blocks overall.
    mn = tmp_path / "mn.docx"
    mx = tmp_path / "mx.docx"
    generate_doc(str(mn), complexity="minimal", sections=8, seed=1)
    generate_doc(str(mx), complexity="maximum", sections=8, seed=1)
    assert len(Document(str(mx)).paragraphs) > len(Document(str(mn)).paragraphs)


def test_blocks_per_section_override(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(
        str(out), complexity="minimal", sections=4, seed=1, blocks_per_section=6
    )
    doc = Document(str(out))
    # 4 section headings + body; the override forces 6 extra blocks per section,
    # so there are far more paragraphs than the minimal default (1) would give.
    assert len(doc.paragraphs) > 4 * 6


def test_invalid_blocks_per_section_raises() -> None:
    with pytest.raises(ValueError, match="blocks_per_section"):
        DocxComplexityGenerator(Complexity.MINIMAL, sections=2, blocks_per_section=0)


def test_seed_is_deterministic(tmp_path: Path) -> None:
    a, b = tmp_path / "a.docx", tmp_path / "b.docx"
    generate_doc(str(a), complexity="complex", sections=10, seed=42)
    generate_doc(str(b), complexity="complex", sections=10, seed=42)
    assert _texts(a) == _texts(b)


def test_different_seed_differs(tmp_path: Path) -> None:
    a, b = tmp_path / "a.docx", tmp_path / "b.docx"
    generate_doc(str(a), complexity="complex", sections=10, seed=1)
    generate_doc(str(b), complexity="complex", sections=10, seed=2)
    assert _texts(a) != _texts(b)


def test_tables_fit_the_page_width(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="maximum", sections=20, seed=5)
    doc = Document(str(out))
    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    for table in doc.tables:
        total = sum(cell.width or 0 for cell in table.rows[0].cells)
        # Column widths sum to no more than the usable page width (+1pt slack).
        assert total <= usable + 12700


def _texts(path: Path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs]
